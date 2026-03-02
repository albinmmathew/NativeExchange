from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_submission_docx():
    doc = Document()

    # Title
    title = doc.add_heading('PROJECT SUBMISSION REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Project Title Section
    doc.add_heading('Project Title', level=1)
    p = doc.add_paragraph('Digital Platform for Micro-Scale Local Producers and Consumers')
    p.runs[0].bold = True

    doc.add_paragraph('---')

    # Introduction Section
    doc.add_heading('Introduction', level=1)
    intro_txt = (
        "The Digital Platform for Micro-Scale Local Producers and Consumers is a modular desktop application built using Java Swing and Oracle Database. "
        "It is designed to bridge the gap between small-scale local producers (vendors) and customers by providing a secure, role-based platform "
        "for inventory discovery and request processing.\n\n"
        "The application implements a robust Role-Based Access Control (RBAC) system, ensuring that Admins, Vendors, "
        "and Customers each have a specialized interface tailored to their specific needs. Key features include "
        "dynamic product discovery, real-time inventory management, and automated request tracking. By utilizing a "
        "modular architecture with a central Dashboard and specialized JPanels, the code remains maintainable, "
        "scalable, and secure."
    )
    doc.add_paragraph(intro_txt)

    doc.add_paragraph('---')

    # Class Diagram Section
    doc.add_heading('Class Diagram', level=1)
    doc.add_paragraph(
        "Note: The digital version of this document contains an interactive Mermaid diagram. "
        "For the printed submission, please insert a high-quality color screenshot of the diagram below."
    )
    
    # Placeholder for Diagram
    doc.add_paragraph("\n[ INSERT CLASS DIAGRAM SCREENSHOT HERE ]\n").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('---')

    # Technical Summary Section
    doc.add_heading('Technical Summary', level=1)
    tech_summary = [
        "UI: Java Swing (CardLayout, BoxLayout)",
        "Database: Oracle 21c (FREEPDB1)",
        "Security: Prepared Statements & RBAC",
        "Architecture: Modular Panel-based Design"
    ]
    for item in tech_summary:
        doc.add_paragraph(item, style='List Bullet')

    # Save the document
    output_path = r'C:\Users\Albin\.gemini\antigravity\brain\72a7d4be-b63d-4af9-bccd-631c52090ce0\submission_document.docx'
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_submission_docx()
